"""监管文档解析模块 - 支持 PDF、DOC、DOCX 格式，提取条文结构"""

import logging
import re
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Generator, Optional

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

logger = logging.getLogger(__name__)

# 条文编号正则：第X条、第X款、第X项等
ARTICLE_PATTERN = re.compile(r"第[一二三四五六七八九十百千零〇]+条")
ARTICLE_ITEM_PATTERN = re.compile(r"第[一二三四五六七八九十百千零〇]+[条款项]")


def _extract_text_from_pdf(file_path: str) -> str:
    """解析 PDF 文件，提取文本内容（处理中文编码）"""
    if PyPDF2 is None:
        raise ImportError("请安装 PyPDF2: pip install PyPDF2")

    text_parts: list[str] = []
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for i, page in enumerate(reader.pages):
                try:
                    content = page.extract_text()
                    if content:
                        if isinstance(content, bytes):
                            content = content.decode("utf-8", errors="replace")
                        text_parts.append(content)
                except Exception as e:
                    logger.warning("PDF 第 %d 页提取失败: %s", i + 1, e)
    except Exception as e:
        logger.error("PDF 读取失败 %s: %s", file_path, e)
        raise

    return "\n\n".join(text_parts)


def _convert_doc_to_docx(doc_path: str) -> Optional[str]:
    """尝试使用 LibreOffice 将 .doc 转为 .docx"""
    doc_path = Path(doc_path).resolve()
    if not doc_path.exists():
        return None

    with tempfile.TemporaryDirectory() as tmpdir:
        out_dir = Path(tmpdir)
        try:
            result = subprocess.run(
                [
                    "soffice",
                    "--headless",
                    "--convert-to",
                    "docx",
                    "--outdir",
                    str(out_dir),
                    str(doc_path),
                ],
                capture_output=True,
                text=True,
                timeout=30,
            )
            if result.returncode != 0:
                logger.warning("LibreOffice 转换失败: %s", result.stderr)
                return None

            docx_path = out_dir / (doc_path.stem + ".docx")
            if docx_path.exists():
                import shutil
                persistent = Path(tempfile.gettempdir()) / f"_docx_{doc_path.stem}.docx"
                shutil.copy(docx_path, persistent)
                return str(persistent)
        except FileNotFoundError:
            logger.warning("未找到 LibreOffice (soffice)，无法解析 .doc 文件")
        except subprocess.TimeoutExpired:
            logger.warning("LibreOffice 转换超时")
        except Exception as e:
            logger.warning("DOC 转换异常: %s", e)

    return None


def _extract_text_from_docx(file_path: str) -> str:
    """解析 DOCX 文件，提取文本内容"""
    if DocxDocument is None:
        raise ImportError("请安装 python-docx: pip install python-docx")

    try:
        doc = DocxDocument(file_path)
        parts: list[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                parts.append(t)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = cell.text.strip()
                    if t:
                        parts.append(t)
        return "\n\n".join(parts)
    except Exception as e:
        logger.error("DOCX 读取失败 %s: %s", file_path, e)
        raise


def _extract_text_from_doc(file_path: str) -> str:
    """解析 .doc 文件：先尝试 python-docx（部分 .doc 实为 OOXML），否则用 LibreOffice 转换"""
    try:
        return _extract_text_from_docx(file_path)
    except Exception:
        pass
    converted = _convert_doc_to_docx(file_path)
    if converted is None:
        raise ValueError(
            f"无法解析 .doc 文件 {file_path}。"
            "请安装 LibreOffice 或将文件手动转换为 .docx 后重试。"
        )
    try:
        return _extract_text_from_docx(converted)
    finally:
        try:
            Path(converted).unlink(missing_ok=True)
        except Exception:
            pass


def _extract_raw_text(file_path: str) -> str:
    """根据扩展名提取原始文本"""
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return _extract_text_from_pdf(file_path)
    if ext == ".docx":
        return _extract_text_from_docx(file_path)
    if ext == ".doc":
        return _extract_text_from_doc(file_path)
    raise ValueError(f"不支持的文件格式: {ext}")


def _split_into_articles(text: str, document_name: str) -> list[dict[str, Any]]:
    """按条文分割文本，识别「第X条」「第X款」等结构"""
    articles: list[dict[str, Any]] = []
    lines = text.replace("\r\n", "\n").split("\n")
    current_id = ""
    current_content: list[str] = []
    buffer_before_first: list[str] = []

    def flush_current():
        nonlocal current_id, current_content
        if current_content:
            content = "\n".join(current_content).strip()
            if content:
                articles.append({
                    "article_id": current_id or "前言/总则",
                    "content": content,
                    "document": document_name,
                })
        current_content = []

    for line in lines:
        line = line.strip()
        if not line:
            if current_content:
                current_content.append("")
            continue

        match = ARTICLE_PATTERN.search(line)
        if match:
            flush_current()
            current_id = match.group(0)
            rest = line[match.end():].strip()
            if rest:
                current_content = [rest]
            else:
                current_content = []
        else:
            sub_match = ARTICLE_ITEM_PATTERN.search(line)
            if sub_match and current_id:
                current_content.append(line)
            elif current_id or articles:
                current_content.append(line)
            else:
                buffer_before_first.append(line)

    flush_current()

    if buffer_before_first and not any(a["article_id"] == "前言/总则" for a in articles):
        intro = "\n".join(buffer_before_first).strip()
        if intro and len(intro) > 20:
            articles.insert(0, {
                "article_id": "前言/总则",
                "content": intro,
                "document": document_name,
            })

    return articles


def parse_document(file_path: str) -> list[dict[str, Any]]:
    """
    解析单个文档，提取条文结构。

    Args:
        file_path: 文档路径（支持 .pdf, .doc, .docx）

    Returns:
        [{"article_id": "第X条", "content": "...", "document": "文档名"}, ...]
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    document_name = path.stem
    logger.info("开始解析文档: %s", document_name)

    try:
        raw_text = _extract_raw_text(file_path)
    except Exception as e:
        logger.error("文档读取失败 %s: %s", file_path, e)
        raise

    if not raw_text or not raw_text.strip():
        logger.warning("文档内容为空: %s", file_path)
        return []

    articles = _split_into_articles(raw_text, document_name)
    logger.info("解析完成 %s: 共 %d 条", document_name, len(articles))
    return articles


def parse_documents(file_paths: list[str]) -> list[dict[str, Any]]:
    """批量解析多个文档，解析失败的会记录日志并跳过"""
    all_articles: list[dict[str, Any]] = []
    for fp in file_paths:
        try:
            articles = parse_document(fp)
            all_articles.extend(articles)
        except Exception as e:
            logger.exception("跳过文档 %s: %s", fp, e)
    return all_articles


# 兼容旧接口
def parse_pdf(file_path: str) -> str:
    """解析 PDF，返回原始文本（兼容）"""
    return _extract_text_from_pdf(file_path)


def parse_docx(file_path: str) -> str:
    """解析 DOCX，返回原始文本（兼容）"""
    return _extract_text_from_docx(file_path)


def parse_txt(file_path: str) -> str:
    """解析 TXT"""
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def load_documents_from_dir(dir_path: str) -> Generator[tuple[str, str], None, None]:
    """从目录加载所有支持的文档，yield (路径, 原始文本)"""
    supported = {".pdf", ".docx", ".doc", ".txt"}
    path = Path(dir_path)
    if not path.exists() or not path.is_dir():
        return
    for f in sorted(path.iterdir()):
        if f.suffix.lower() in supported:
            try:
                raw = _extract_raw_text(str(f))
                if raw.strip():
                    yield str(f), raw
            except Exception as e:
                logger.warning("解析 %s 失败: %s", f, e)
